import os
from pathlib import Path
from azure.core.credentials import AzureKeyCredential
from azure.ai.formrecognizer import DocumentAnalysisClient
import docx
from typing import Optional
from dotenv import load_dotenv

from logger import CustomLogger

log = CustomLogger("DocumentProcessor", log_file="document_processor.log")

load_dotenv()

class UnsupportedFileTypeError(Exception):
    pass

def get_file_type(file_path: str) -> str:
    """Determine file type from extension."""
    return Path(file_path).suffix.lower()

def read_txt_file(file_path: str) -> str:
    """Read content from a text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try different encodings if UTF-8 fails
        encodings = ['latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError("Failed to decode text file with multiple encodings")

def read_docx_file(file_path: str) -> str:
    """Extract text from a DOCX file."""
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    
    return '\n'.join(full_text)

def read_pdf_with_azure(file_path: str, endpoint: Optional[str] = None, key: Optional[str] = None) -> str:
    """Extract text from PDF using Azure Document Intelligence."""
    try:
        # Use provided credentials or fall back to environment variables
        endpoint = endpoint or os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
        key = key or os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY")
        
        if not endpoint or not key:
            raise ValueError("Azure Document Intelligence credentials not configured")
        
        document_analysis_client = DocumentAnalysisClient(
            endpoint=endpoint, credential=AzureKeyCredential(key)
        )

        with open(file_path, "rb") as file:
            document_bytes = file.read()

        poller = document_analysis_client.begin_analyze_document(
            "prebuilt-read", document_bytes
        )
        result = poller.result()
        return result.content
    except Exception as e:
        log.log_error(f"Azure PDF processing failed: {str(e)}")
        raise

async def analyze_document(file_path: str) -> str:
    """
    Analyze document from various file formats.
    Supports PDF, DOC, DOCX, and TXT files.
    
    Args:
        file_path: Path to the local file
    Returns:
        Extracted text content from the document
    """
    try:
        # Verify file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_path = str(Path(file_path))  # Normalize path
        file_type = get_file_type(file_path)
        
        if file_type not in ['.pdf', '.doc', '.docx', '.txt']:
            raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")

        log.log_info(f"Processing {file_type} file: {file_path}")

        # Process based on file type
        if file_type == '.txt':
            log.log_info("Processing TXT file")
            full_text = read_txt_file(file_path)
            
        elif file_type in ['.doc', '.docx']:
            log.log_info("Processing DOC/DOCX file")
            if file_type == '.doc':
                # For now, we'll raise an error for .doc files
                # You might want to add doc to docx conversion here
                raise UnsupportedFileTypeError("DOC format is not supported, please convert to DOCX")
            full_text = read_docx_file(file_path)
            
        elif file_type == '.pdf':
            log.log_info("Processing PDF file using Azure Document Intelligence")
            full_text = read_pdf_with_azure(file_path)

        # Post-processing
        if not full_text:
            raise ValueError("No text content extracted from document")

        # Remove excessive whitespace and normalize line endings
        full_text = '\n'.join(line.strip() for line in full_text.splitlines() if line.strip())
        log.log_info(f"Successfully extracted {len(full_text)} characters from {file_type} file")
        
        return full_text

    except UnsupportedFileTypeError as e:
        log.log_error(f"Unsupported file type error: {str(e)}")
        raise

    except Exception as e:
        log.log_error(f"Error analyzing document: {str(e)}")
        raise

def validate_file_type(file_path: str) -> bool:
    """
    Validate if the file type is supported.
    Returns True if supported, False otherwise.
    """
    allowed_extensions = {'.pdf', '.doc', '.docx', '.txt'}
    return get_file_type(file_path) in allowed_extensions

async def get_document_metadata(file_path: str) -> dict:
    """
    Get metadata about the local document.
    Returns a dictionary with file type, size, and other relevant information.
    """
    try:
        file_stat = os.stat(file_path)
        metadata = {
            "file_type": get_file_type(file_path),
            "size": file_stat.st_size,
            "created_at": file_stat.st_ctime,
            "updated_at": file_stat.st_mtime,
            "path": str(Path(file_path).absolute())
        }
        return metadata
    except Exception as e:
        log.log_error(f"Error getting document metadata: {str(e)}")
        raise